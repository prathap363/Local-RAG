import logging
from pathlib import Path
from threading import Lock
from typing import Any
import unicodedata

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from langchain_experimental.text_splitter import SemanticChunker
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_postgres import PGVector
from pydantic import BaseModel
from pydantic_settings import BaseSettings, SettingsConfigDict


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("rag-backend")


class Settings(BaseSettings):
    """Centralized runtime settings loaded from environment variables.

    This class maps environment variables to typed Python attributes so the
    application can rely on validated, strongly-typed configuration values.
    Values are read once at startup from process env vars and optional `.env`
    files.
    """

    model_config = SettingsConfigDict(
        # Load from .env when present (useful for local runs), while still
        # allowing real environment variables to override these values.
        env_file=(".env", "../.env"),
        env_file_encoding="utf-8",
        case_sensitive=False,
        # Allow fields that start with "model_" (for model runner settings).
        protected_namespaces=("settings_",),
    )

    postgres_host: str
    postgres_port: int
    postgres_db: str
    postgres_user: str
    postgres_password: str

    collection_name: str
    data_dir: str
    chunk_size: int
    chunk_overlap: int
    chunking_mode: str = "recursive"
    top_k: int
    reset_collection_on_start: bool

    model_runner_base_url: str
    model_runner_api_key: str
    embedding_model: str
    query_model: str
    answer_model: str

    cors_allow_origins: str

    @property
    def connection_string(self) -> str:
        """Build SQLAlchemy-compatible connection string for PGVector."""
        return (
            f"postgresql+psycopg://{self.postgres_user}:{self.postgres_password}"
            f"@{self.postgres_host}:{self.postgres_port}/{self.postgres_db}"
        )


class ChatRequest(BaseModel):
    """Request payload accepted by `/chat`.

    Fields:
    - query: end-user natural language question.
    """

    query: str


class ChatResponse(BaseModel):
    """Structured response returned by `/chat`.

    Fields:
    - answer: model-generated answer text.
    - augmented_query: rewritten query used for vector search.
    - sources: unique list of source files used in retrieved context.
    """

    answer: str
    augmented_query: str
    sources: list[str]


# Application settings and FastAPI app initialization happen once at import time.
settings = Settings()
app = FastAPI(title="Simple RAG Backend", version="1.0.0")


def _mask(value: str, reveal: int = 2) -> str:
    """Mask sensitive values before writing them to logs.

    Example:
    - input: `abcdef`, reveal=2
    - output: `****ef`
    """
    if not value:
        return ""
    if len(value) <= reveal:
        return "*" * len(value)
    return ("*" * (len(value) - reveal)) + value[-reveal:]


def log_loaded_settings() -> None:
    """Log effective runtime configuration loaded from env/.env.

    Sensitive values (passwords/API keys) are masked before logging.
    This helps with startup diagnostics without exposing secrets.
    """
    logger.info("Loaded settings:")
    logger.info("  POSTGRES_HOST=%s", settings.postgres_host)
    logger.info("  POSTGRES_PORT=%s", settings.postgres_port)
    logger.info("  POSTGRES_DB=%s", settings.postgres_db)
    logger.info("  POSTGRES_USER=%s", settings.postgres_user)
    logger.info("  POSTGRES_PASSWORD=%s", _mask(settings.postgres_password))
    logger.info("  COLLECTION_NAME=%s", settings.collection_name)
    logger.info("  DATA_DIR=%s", settings.data_dir)
    logger.info("  CHUNK_SIZE=%s", settings.chunk_size)
    logger.info("  CHUNK_OVERLAP=%s", settings.chunk_overlap)
    logger.info("  CHUNKING_MODE=%s", settings.chunking_mode)
    logger.info("  TOP_K=%s", settings.top_k)
    logger.info("  RESET_COLLECTION_ON_START=%s", settings.reset_collection_on_start)
    logger.info("  MODEL_RUNNER_BASE_URL=%s", settings.model_runner_base_url)
    logger.info("  MODEL_RUNNER_API_KEY=%s", _mask(settings.model_runner_api_key))
    logger.info("  EMBEDDING_MODEL=%s", settings.embedding_model)
    logger.info("  QUERY_MODEL=%s", settings.query_model)
    logger.info("  ANSWER_MODEL=%s", settings.answer_model)
    logger.info("  CORS_ALLOW_ORIGINS=%s", settings.cors_allow_origins)

# Allow one or more CORS origins via comma-separated env var.
origins = [o.strip() for o in settings.cors_allow_origins.split(",") if o.strip()]
if not origins:
    origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_state_lock = Lock()
_initialized = False
_init_error: str | None = None

_embeddings: OpenAIEmbeddings | None = None
_query_llm: ChatOpenAI | None = None
_answer_llm: ChatOpenAI | None = None
_vector_store: PGVector | None = None


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".py",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".log",
    ".rst",
    ".pdf",
    ".docx",
}


def _read_pdf(path: Path) -> str:
    """Extract text from all pages of a PDF file.

    Returns one concatenated string that preserves page order.
    """

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return "\n".join(chunks).strip()


def _read_docx(path: Path) -> str:
    """Extract text from a DOCX file, preserving paragraph order."""

    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs).strip()


def _sanitize_text(text: str) -> str:
    """Normalize and clean text before embedding.

    Why this exists:
    - Some embedding backends reject specific control/surrogate characters.
    - Normalization makes text representation more consistent.

    Behavior:
    - Apply Unicode NFKC normalization.
    - Keep newline/carriage-return/tab.
    - Drop other Unicode category `C*` control-like chars.
    - Trim each line.
    """

    normalized = unicodedata.normalize("NFKC", text)
    cleaned_chars: list[str] = []
    # Character-by-character clean-up keeps line structure but removes
    # problematic control characters for downstream tokenization.
    for ch in normalized:
        if ch in {"\n", "\r", "\t"}:
            cleaned_chars.append(ch)
            continue

        # Drop control/format/surrogate/private-use codepoints that can trigger
        # tokenizer failures in some local embedding backends.
        if unicodedata.category(ch)[0] == "C":
            continue

        cleaned_chars.append(ch)

    cleaned = "".join(cleaned_chars)
    # Keep line boundaries but trim noisy whitespace per line.
    return "\n".join(line.strip() for line in cleaned.splitlines())


def _ascii_fallback_text(text: str) -> str:
    """Reduce text to ASCII-safe content for strict embedding backends.

    This fallback is used when a chunk fails to embed due to tokenizer errors.
    """

    ascii_text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    # Collapse excessive whitespace while preserving paragraph breaks.
    lines = [" ".join(line.split()) for line in ascii_text.splitlines()]
    return "\n".join(line for line in lines if line)


def load_documents(directory: str) -> list[Document]:
    """Load supported files recursively from `directory` into Documents.

    File handling strategy:
    - `.pdf` -> page text extraction
    - `.docx` -> paragraph extraction
    - others -> UTF-8 text read with ignore errors

    Each loaded item carries metadata with relative and absolute paths.
    """

    root = Path(directory)
    if not root.exists() or not root.is_dir():
        raise ValueError(f"Data directory does not exist or is not a folder: {directory}")

    documents: list[Document] = []
    # Walk every file recursively under the configured corpus root.
    for file_path in root.rglob("*"):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:
            # Route by extension so each format uses an appropriate extractor.
            if file_path.suffix.lower() == ".pdf":
                content = _read_pdf(file_path)
            elif file_path.suffix.lower() == ".docx":
                content = _read_docx(file_path)
            else:
                content = file_path.read_text(encoding="utf-8", errors="ignore")
            content = _sanitize_text(content)
        except Exception as exc:
            logger.warning("Skipping unreadable file %s: %s", file_path, exc)
            continue

        if not content.strip():
            continue

        # Keep both a relative source and absolute path for easy tracing.
        rel = str(file_path.relative_to(root))
        documents.append(
            Document(
                page_content=content,
                metadata={"source": rel, "absolute_path": str(file_path)},
            )
        )

    return documents


def split_documents(documents: list[Document]) -> list[Document]:
    """Split documents based on configured chunking mode.

    Supported values for `CHUNKING_MODE`:
    - `recursive` (default): fixed-size overlapping chunks
    - `semantic`: embedding-aware semantic boundaries

    Any invalid mode or semantic chunking failure falls back to recursive.
    """

    mode = settings.chunking_mode.strip().lower()
    if mode not in {"recursive", "semantic"}:
        logger.warning(
            "Unsupported CHUNKING_MODE '%s'; falling back to recursive",
            settings.chunking_mode,
        )
        mode = "recursive"

    if mode == "semantic":
        if _embeddings is None:
            logger.warning("Embeddings not initialized for semantic chunking; using recursive")
        else:
            try:
                semantic_splitter = SemanticChunker(embeddings=_embeddings)
                semantic_chunks = semantic_splitter.split_documents(documents)
                if semantic_chunks:
                    return semantic_chunks
                logger.warning("Semantic chunking returned no chunks; falling back to recursive")
            except Exception as exc:
                logger.warning("Semantic chunking failed; falling back to recursive: %s", exc)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    return splitter.split_documents(documents)


def _add_documents_resilient(chunks: list[Document], batch_size: int = 16) -> tuple[int, int]:
    """Add chunks to PGVector with fallback handling for embed failures.

    Strategy:
    1) Try batch insert for throughput.
    2) If batch fails, retry each chunk individually.
    3) If individual chunk fails due to invalid tokens, retry with ASCII fallback.
    4) Track added vs skipped counts for observability.
    """

    if _vector_store is None:
        raise RuntimeError("Vector store is not initialized")

    added = 0
    skipped = 0

    # Process in batches for efficiency while retaining granular recovery paths.
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start : start + batch_size]
        try:
            _vector_store.add_documents(batch)
            added += len(batch)
            continue
        except Exception as exc:
            logger.warning(
                "Batch add failed for chunks %s-%s; retrying one by one. Error: %s",
                start,
                start + len(batch) - 1,
                exc,
            )

        # Isolate problematic chunks so one bad chunk does not fail whole ingest.
        for index, chunk in enumerate(batch):
            try:
                _vector_store.add_documents([chunk])
                added += 1
            except Exception as exc:
                source = chunk.metadata.get("source", "unknown")
                err_text = str(exc)
                if "invalid tokens" in err_text.lower():
                    try:
                        fallback_chunk = Document(
                            page_content=_ascii_fallback_text(chunk.page_content),
                            metadata={**chunk.metadata, "sanitized": "ascii-fallback"},
                        )
                        if fallback_chunk.page_content.strip():
                            _vector_store.add_documents([fallback_chunk])
                            added += 1
                            logger.warning(
                                "Applied ASCII fallback for chunk at global index %s from '%s'",
                                start + index,
                                source,
                            )
                            continue
                    except Exception as fallback_exc:
                        logger.warning(
                            "ASCII fallback failed for chunk at global index %s from '%s': %s",
                            start + index,
                            source,
                            fallback_exc,
                        )

                skipped += 1
                logger.warning(
                    "Skipping chunk at global index %s from '%s': %s",
                    start + index,
                    source,
                    exc,
                )

    return added, skipped


def _ensure_initialized(force_reingest: bool = False) -> None:
    """Initialize runtime dependencies and ingest corpus when needed.

    This method is idempotent under a lock:
    - First call initializes embeddings/LLMs/vector store and ingests data.
    - Later calls return quickly unless `force_reingest=True`.
    """

    global _initialized, _init_error
    global _embeddings, _query_llm, _answer_llm, _vector_store

    # Lock protects shared globals from concurrent startup/chat/ingest access.
    with _state_lock:
        if _initialized and not force_reingest:
            return

        _init_error = None

        # All model calls are routed through an OpenAI-compatible local endpoint.
        _embeddings = OpenAIEmbeddings(
            model=settings.embedding_model,
            base_url=settings.model_runner_base_url,
            api_key=settings.model_runner_api_key,
        )
        _query_llm = ChatOpenAI(
            model=settings.query_model,
            base_url=settings.model_runner_base_url,
            api_key=settings.model_runner_api_key,
            temperature=0,
        )
        _answer_llm = ChatOpenAI(
            model=settings.answer_model,
            base_url=settings.model_runner_base_url,
            api_key=settings.model_runner_api_key,
            temperature=0.2,
        )

        _vector_store = PGVector(
            embeddings=_embeddings,
            collection_name=settings.collection_name,
            connection=settings.connection_string,
            use_jsonb=True,
            # Drop/recreate collection only during startup-style initialization.
            pre_delete_collection=settings.reset_collection_on_start and not _initialized,
        )

        docs = load_documents(settings.data_dir)
        if not docs:
            logger.warning("No supported files found under %s", settings.data_dir)
        else:
            chunks = split_documents(docs)
            logger.info(
                "Loaded %s documents and generated %s chunks", len(docs), len(chunks)
            )
            added, skipped = _add_documents_resilient(chunks)
            logger.info(
                "Ingested %s chunks (skipped %s) into pgvector collection '%s'",
                added,
                skipped,
                settings.collection_name,
            )

        _initialized = True


def _expand_query(user_query: str) -> str:
    """Rewrite user query for improved semantic retrieval recall.

    If rewriting fails or returns empty content, the original query is used.
    """

    if _query_llm is None:
        return user_query

    prompt = (
        "Rewrite the user question as an optimized search query for document retrieval. "
        "Keep original intent and key entities. Return only the rewritten query.\n\n"
        f"User question: {user_query}"
    )
    result = _query_llm.invoke(prompt)
    content = getattr(result, "content", "")
    if isinstance(content, list):
        content = " ".join(str(x) for x in content)
    rewritten = str(content).strip().strip('"').strip("'")
    return rewritten or user_query


def _merge_retrieval_results(primary: list[Document], secondary: list[Document], k: int) -> list[Document]:
    """Merge two ranked retrieval lists and remove near-duplicate entries.

    Dedup key uses `(source, first_200_chars)` as a lightweight stable identity.
    This preserves ranking while avoiding repeated chunks in answer context.
    """

    merged: list[Document] = []
    seen: set[tuple[str, str]] = set()

    # Keep ordering preference: `primary` results first, then `secondary` fill-ins.
    for doc in primary + secondary:
        key = (
            str(doc.metadata.get("source", "unknown")),
            doc.page_content[:200],
        )
        if key in seen:
            continue
        seen.add(key)
        merged.append(doc)
        if len(merged) >= k:
            break

    return merged


def _answer_question(user_query: str, augmented_query: str, docs: list[Document]) -> str:
    """Generate final answer grounded on retrieved document chunks.

    The prompt explicitly includes both the original and augmented query plus
    the retrieved context block to guide grounded generation.
    """

    if _answer_llm is None:
        raise HTTPException(status_code=500, detail="Answer model is not initialized")

    context = "\n\n".join(doc.page_content for doc in docs)
    prompt = (
        "You are a helpful assistant answering strictly from the provided context. "
        "If the context is insufficient, say so clearly.\n\n"
        f"User question: {user_query}\n"
        f"Augmented retrieval query: {augmented_query}\n\n"
        f"Context:\n{context}"
    )
    result = _answer_llm.invoke(prompt)
    content = getattr(result, "content", "")
    if isinstance(content, list):
        content = " ".join(str(x) for x in content)
    return str(content).strip()


@app.on_event("startup")
def startup_ingest() -> None:
    """FastAPI startup hook: initialize runtime and run initial ingestion."""

    global _init_error
    try:
        log_loaded_settings()
        _ensure_initialized(force_reingest=False)
    except Exception as exc:
        _init_error = str(exc)
        logger.exception("Startup initialization failed")


@app.get("/health")
def health() -> dict[str, Any]:
    """Health endpoint returning service and initialization state."""

    return {
        "status": "ok",
        "initialized": _initialized,
        "init_error": _init_error,
        "collection": settings.collection_name,
    }


@app.post("/ingest")
def ingest() -> dict[str, Any]:
    """Manual endpoint to force re-ingestion from `DATA_DIR`.

    Useful after adding/modifying files in the mounted corpus directory.
    """

    global _initialized
    try:
        # Force reingestion while preserving existing collection setting behavior.
        _initialized = False
        _ensure_initialized(force_reingest=True)
        return {"status": "ok", "message": "Ingestion completed"}
    except Exception as exc:
        logger.exception("Ingestion failed")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/chat", response_model=ChatResponse)
def chat(payload: ChatRequest) -> ChatResponse:
    """RAG chat endpoint: expand query, retrieve context, generate answer.

    Runtime flow:
    1) Ensure runtime is initialized.
    2) Build augmented query.
    3) Retrieve top-k from original and augmented queries.
    4) Merge/dedup retrieval outputs.
    5) Generate grounded answer and return sources.
    """

    try:
        _ensure_initialized(force_reingest=False)
    except Exception as exc:
        logger.exception("Initialization failed")
        raise HTTPException(status_code=500, detail=f"Initialization failed: {exc}") from exc

    if _vector_store is None:
        raise HTTPException(status_code=500, detail="Vector store is not initialized")

    augmented_query = _expand_query(payload.query)

    # Retrieve using both original and rewritten query to avoid rewrite misses.
    retriever = _vector_store.as_retriever(search_kwargs={"k": settings.top_k})
    docs_original = retriever.invoke(payload.query)
    docs_augmented = retriever.invoke(augmented_query)
    docs = _merge_retrieval_results(docs_original, docs_augmented, settings.top_k)

    answer = _answer_question(payload.query, augmented_query, docs)

    # Preserve unique source file list in response metadata.
    sources: list[str] = []
    for doc in docs:
        source = str(doc.metadata.get("source", "unknown"))
        if source not in sources:
            sources.append(source)

    return ChatResponse(answer=answer, augmented_query=augmented_query, sources=sources)


@app.get("/")
def root() -> dict[str, str]:
    """Basic root endpoint for quick service sanity check."""

    return {"message": "Simple RAG backend is running"}
